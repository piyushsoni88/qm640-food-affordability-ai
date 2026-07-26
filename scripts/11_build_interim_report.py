from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\piyus\Downloads\QM 640 Interim Report template-1 (1).docx")
OUTPUT = ROOT / "reports" / "Piyush_Soni_QM640_Interim_Report.docx"
FIGURES = ROOT / "reports" / "figures"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
GRAY = "666666"


def style_by_name(doc: Document, name: str):
    for style in doc.styles:
        if style.name == name:
            return style
    raise KeyError(f"Style not found: {name}")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=12, bold=None, italic=None, color=None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in [fld_begin, instr, fld_separate, display, fld_end]:
        run._r.append(node)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    normal = style_by_name(doc, "Normal")
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name, size, centered in [
        ("Heading 1", 12, True),
        ("Heading 2", 12, False),
        ("Heading 3", 12, False),
    ]:
        style = style_by_name(doc, style_name)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        )


def configure_section(section, title_page=False) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = title_page


def configure_header(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Interim Report")
    set_run_font(run, size=9, color=GRAY)
    paragraph.add_run("\t")
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.45), WD_ALIGN_PARAGRAPH.RIGHT
    )
    add_field(paragraph, "PAGE")


def add_paragraph(doc, text: str, bold_lead: str | None = None, no_indent=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Inches(0 if no_indent else 0.5)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc, items):
    for index, item in enumerate(items, 1):
        paragraph = doc.add_paragraph(style=style_by_name(doc, "List Paragraph"))
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph.paragraph_format.space_after = Pt(0)
        colon = item.find(":")
        if 0 < colon < 28:
            label = paragraph.add_run(item[: colon + 1] + " ")
            set_run_font(label, bold=True)
            item = item[colon + 1 :].strip()
        else:
            label = paragraph.add_run(f"Action {index}. ")
            set_run_font(label, bold=True)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=style_by_name(doc, f"Heading {level}"))
    for run in paragraph.runs:
        set_run_font(run, bold=True)
    return paragraph


def add_caption(doc, label: str, title: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=10, bold=True)
    title_run = paragraph.add_run(f"\n{title}")
    set_run_font(title_run, size=10, italic=True)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = style_by_name(doc, "TableNormal")
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        cell.text = str(text)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=font_size, bold=True)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.text = str(value)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Inches(0)
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=font_size)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc, file_name: str, label: str, title: str, width=6.25):
    add_caption(doc, label, title)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    run.add_picture(str(FIGURES / file_name), width=Inches(width))


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    clear_document_body(doc)
    configure_styles(doc)
    first = doc.sections[0]
    configure_section(first, title_page=True)

    # Title page
    for _ in range(4):
        doc.add_paragraph()
    title_lines = [
        ("Data Analytics Capstone", True),
        (
            "Forecasting Essential Food Price Shocks and Household Affordability Stress in India",
            True,
        ),
        (
            "An Explainable AI Decision-Intelligence Framework Using Public Market, Agricultural, Climatic, and Economic Data",
            False,
        ),
        ("Interim Report", True),
        ("Piyush Soni", False),
        ("Walsh College", False),
        ("QM640: Data Analytics Capstone", False),
        ("Mentor: Prof. Rishabh Pandey", False),
        ("Summer 2026 Term", False),
        ("July 27, 2026", False),
    ]
    for index, (text, bold) in enumerate(title_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 2
        p.paragraph_format.space_after = Pt(0 if index else 12)
        run = p.add_run(text)
        set_run_font(run, bold=bold)

    report_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(report_section)
    configure_header(report_section)

    add_heading(doc, "GitHub Repository and Data Availability", 1)
    p = add_paragraph(
        doc,
        "The public repository containing the acquisition code, curated data, analysis outputs, tests, and documentation is:",
        no_indent=True,
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(
        p,
        "github.com/piyushsoni88/qm640-food-affordability-ai",
        "https://github.com/piyushsoni88/qm640-food-affordability-ai",
    )
    add_paragraph(
        doc,
        "GitHub-ready curated datasets are stored under data/curated/. Raw bulk archives are retained locally under data/raw/ and are reproducibly reacquired by scripts because the 51.5 MB source layer expands to millions of global records and should not be duplicated unnecessarily in routine version control. The source manifest records URLs, timestamps, file sizes, SHA-256 checksums, row counts, and derived-file paths.",
    )
    add_caption(doc, "Table 1", "Interim data snapshot")
    add_table(
        doc,
        ["Data product", "Rows", "Coverage", "Repository path"],
        [
            ["FAOSTAT India food CPI", "924", "2000-2025", "data/curated/faostat_india_consumer_price_indices.csv.gz"],
            ["FAOSTAT India producer prices", "5,665", "1961-2024", "data/curated/faostat_india_producer_prices.csv.gz"],
            ["FAOSTAT India production", "26,432", "1961-2024", "data/curated/faostat_india_crop_production.csv.gz"],
            ["NASA POWER climate", "3,780", "2005-2025; 15 regions", "data/curated/nasa_power_india_15_regions_monthly_2005_2025.csv.gz"],
            ["World Bank benchmarks", "792", "1960-2025", "data/curated/world_bank_pink_sheet_food_energy_monthly.csv.gz"],
            ["Derived analytical panel", "30,240", "2005-2025; 15 x 8", "data/curated/india_food_affordability_panel_15x8_2005_2025.csv.gz"],
        ],
        widths=[1.7, 0.7, 1.4, 2.7],
        font_size=8,
    )
    add_paragraph(
        doc,
        "Reproduction command: python scripts/09_build_interim_research_dataset.py --start-year 2005 --end-year 2025, followed by python scripts/10_run_interim_analysis.py. The data.gov.in AGMARKNET endpoint requires a personal API key; therefore, its collector is included and tested for configuration, but no unauthorized key or fabricated mandi history is included.",
    )

    add_heading(doc, "Introduction", 1)
    add_heading(doc, "Background and Context", 2)
    add_paragraph(
        doc,
        "Food-price volatility is simultaneously a household-welfare problem, a procurement risk, and a policy-monitoring challenge. Essential products such as rice, wheat, onion, potato, tomato, pulses, edible oil, and sugar have different production cycles, storage lives, market structures, and exposure to rainfall and transport costs. Their prices can respond to crop losses, market-arrival contractions, fuel costs, trade restrictions, festival demand, and speculative behavior. When signals are detected only after retail prices rise, households lose purchasing power and food businesses face margin compression, stock-outs, or emergency purchases.",
    )
    add_paragraph(
        doc,
        "The welfare effect is unequal because food absorbs a larger share of expenditure among lower-income households. The Household Consumption Expenditure Survey 2022-2023 reported food shares of approximately 46.3% of rural Monthly Per Capita Consumption Expenditure and 39.1% of urban expenditure (Ministry of Statistics and Programme Implementation [MOSPI], 2024). A useful system must therefore do more than forecast a price index. It must disclose uncertainty, explain the variables associated with predicted change, and translate pressure on essential-food costs into a segment-sensitive affordability measure.",
    )
    add_paragraph(
        doc,
        "India has rich but fragmented public data. Government and multilateral sources publish consumer prices, producer prices, crop production, household expenditure, market prices, rainfall, and international commodity benchmarks. FAOSTAT provides standardized food CPI, producer-price, and production series; NASA POWER provides reproducible regional climate histories; the World Bank Pink Sheet supplies monthly global commodity benchmarks; DES publishes agricultural price compendia; and AGMARKNET provides the intended mandi-level backbone. Their combination supports a staged research design in which a reliable national prototype is completed before the key-dependent regional market backfill.",
    )
    add_heading(doc, "Problem Statement", 2)
    add_paragraph(
        doc,
        "Indian households and food-sector enterprises face recurring uncertainty because essential-food prices are volatile and respond to interacting climatic, agricultural, seasonal, market, and macroeconomic forces. Although public data exist, they are distributed across incompatible formats and are often used only for retrospective monitoring. The current gap is an integrated and explainable early-warning framework that predicts one- to three-month food-price movement, classifies abnormal shocks, quantifies model uncertainty, and translates predicted cost pressure into a Household Food Affordability Stress Index (HFASI).",
    )
    add_heading(doc, "Purpose of the Study", 2)
    add_paragraph(
        doc,
        "This quantitative longitudinal study aims to develop and validate a reproducible decision-intelligence framework for essential-food price shocks and affordability stress in India. It connects four tasks: identify statistically important drivers, forecast future price levels, classify abnormal upward shocks, and estimate affordability stress when expected food-cost growth exceeds household purchasing-power growth. Transparent baselines are compared with machine-learning models using time-ordered validation, while feature importance and SHapley Additive exPlanations are planned for final interpretability.",
    )
    add_heading(doc, "Interim Project Status (Progress Snapshot)", 2)
    add_bullets(
        doc,
        [
            "Completed: repository scaffold, source registry, reproducible collectors, 51.5 MB local raw download, 37,593 source-specific India, climate, and benchmark rows, a 30,240-row derived panel, checksums, data-quality audit, EDA, baseline forecasting, shock classification, and preliminary statistical inference.",
            "In progress: AGMARKNET historical acquisition using an authorized data.gov.in key, market-name harmonization, HCES expenditure-weight extraction, HFASI calibration, and SHAP-based explanation.",
            "Pending: full rolling-origin evaluation at one-, two-, and three-month horizons; probability-threshold tuning for rare shocks; prediction intervals; decision-loss backtesting; robustness checks; and final recommendations.",
        ],
    )

    add_heading(doc, "Scope and Objectives", 1)
    add_paragraph(
        doc,
        "The intended final unit of analysis is commodity-region-month. The extraction architecture targets 2006-2025 for the market-price backbone and 2011-2025 for confirmatory models requiring official CPI/CFPI covariates. Eight commodities and 15 states or representative regional points are retained at interim stage. The prototype panel contains 8 x 15 x 252 = 30,240 unique keys from January 2005 through December 2025. Region-specific NASA observations are combined with national price and production proxies; those national series are explicitly labeled and are not misrepresented as regional mandi prices.",
    )
    add_caption(doc, "Table 2", "Research questions and interim operationalization")
    add_table(
        doc,
        ["Research question", "Interim operationalization", "Final evidence"],
        [
            ["RQ1: Which variables influence food-price change?", "HAC-robust OLS on next-month food CPI change; lag, global-price, climate, and seasonal predictors.", "Commodity-region panel coefficients, stability, VIF, confidence intervals."],
            ["RQ2: Do ML models outperform statistical baselines?", "Persistence, ridge regression, and random forest compared on a fixed 2022-2025 holdout plus five rolling folds.", "Expanded horizons, additional models, paired loss tests, intervals."],
            ["RQ3: Can forecasts produce a reliable HFASI?", "HFASI formula and HCES food-share weights retained from synopsis; calibration not yet claimed.", "Reliability, criterion validity, rank stability, sensitivity."],
            ["RQ4: Does the framework improve decisions?", "Decision use cases and loss functions specified; no causal benefit claim at interim stage.", "Budget, procurement, inventory, and policy backtests."],
        ],
        widths=[2.0, 2.4, 2.1],
        font_size=8,
    )

    add_heading(doc, "Literature Survey", 1)
    add_heading(doc, "Literature Review Approach", 2)
    add_paragraph(
        doc,
        "The review carried forward the approved synopsis search and added source-method documentation needed for the interim implementation. Sources were selected when they addressed at least one of four themes: agricultural price forecasting, shock detection, model evaluation and interpretability, or household affordability. Peer-reviewed articles and official institutional publications were prioritized. Studies that used random train-test splits without temporal safeguards were treated as methodological context rather than direct performance benchmarks.",
    )
    add_caption(doc, "Table 3", "Literature relevance matrix (part 1)")
    add_table(
        doc,
        ["Source", "Context/data", "Method", "Finding and project relevance"],
        [
            ["Madaan et al. (2019)", "Indian agricultural commodities", "Forecasting and anomaly detection", "Supports joint price forecasting and shock detection for RQ1-RQ2."],
            ["Ma et al. (2019)", "Indian produce prices", "Collaborative filtering and adaptive neighbors", "Shows interpretable local forecasting can support farmers and markets."],
            ["Jain et al. (2020)", "Emerging-economy crop series", "Quality-aware framework", "Motivates explicit coverage, missingness, and series-quality controls."],
            ["Bhardwaj et al. (2023)", "Agricultural crop prices", "Deep learning", "Supports nonlinear candidates only after sufficient market-level history exists."],
            ["Theofilou et al. (2025)", "Staple-crop systematic review", "Systematic synthesis", "Identifies validation, exogenous-data, and generalizability gaps."],
        ],
        widths=[1.2, 1.45, 1.45, 2.4],
        font_size=8,
    )
    add_caption(doc, "Table 4", "Literature relevance matrix (part 2)")
    add_table(
        doc,
        ["Source", "Context/data", "Method", "Finding and project relevance"],
        [
            ["Makridakis et al. (2018)", "Forecasting evidence", "Statistical vs. ML critique", "Justifies strong simple baselines and out-of-sample comparison."],
            ["Hyndman & Koehler (2006)", "Forecast accuracy", "Metric analysis", "Supports MAE, RMSE, and scale-aware metric selection."],
            ["Lundberg & Lee (2017)", "Model interpretation", "SHAP", "Provides the planned explanation method for RQ1 and user trust."],
            ["Malesios et al. (2020)", "Food-price shocks", "Change-point analysis", "Supports abnormal-event definitions beyond arbitrary percentage rules."],
            ["Cattaneo et al. (2023)", "Diet affordability under shocks", "Policy synthesis", "Links price shocks to household affordability and mitigation."],
            ["Akter & Basher (2014)", "Rural household food security", "Shock and welfare analysis", "Supports expenditure-share-sensitive interpretation for RQ3."],
        ],
        widths=[1.2, 1.45, 1.45, 2.4],
        font_size=8,
    )
    add_heading(doc, "Thematic Synthesis", 2)
    add_paragraph(
        doc,
        "First, forecasting research consistently finds that persistence and seasonal structure are difficult to beat, especially when samples are short or regimes change. This makes a naive benchmark a scientific control rather than a trivial competitor. Tree ensembles can capture nonlinearities but cannot extrapolate trends beyond the range learned from training data. The interim results demonstrate this risk: the random forest fit historical structure yet failed badly when the test period contained higher CPI levels.",
    )
    add_paragraph(
        doc,
        "Second, high predictive accuracy is not sufficient for a shock-warning system. Rare-event class imbalance can produce high accuracy while delivering zero recall for the events that matter. Threshold-independent discrimination, calibration, precision-recall analysis, and cost-sensitive decision thresholds are therefore necessary. Third, affordability is a distributional outcome, not merely a price index. A common price shock imposes different stress when food expenditure shares and purchasing-power growth differ across rural, urban, and expenditure groups.",
    )
    add_paragraph(
        doc,
        "Finally, reproducibility and interpretability are central design requirements. Time-ordered validation prevents future information from leaking into model selection; source manifests and checksums preserve provenance; and SHAP or standardized coefficients can communicate drivers. However, explanations remain predictive associations. Climate, prices, and production may share trends or omitted causes, so causal language is avoided unless a separate identification strategy is justified.",
    )

    add_heading(doc, "Data Description", 1)
    add_heading(doc, "Data Sources and Access", 2)
    add_paragraph(
        doc,
        "The interim dataset uses five accessible sources. FAOSTAT bulk downloads provide monthly India food CPI, annual producer prices, and crop/livestock production. NASA POWER supplies monthly precipitation, temperature, and relative humidity for 15 regional points. The World Bank Pink Sheet supplies international rice, wheat, edible-oil, sugar, and crude-oil benchmarks. The Directorate of Economics and Statistics (DES) 2024 agricultural price publication is retained locally as an official validation source. AGMARKNET is the planned primary market-price source, but its API requires an authorized personal key; the report therefore separates observed interim proxies from unavailable historical mandi records.",
    )
    add_heading(doc, "Dataset Overview", 2)
    add_paragraph(
        doc,
        "The national modeling table contains 252 monthly observations from January 2005 through December 2025. After lag construction and next-month target alignment, 192 observations from January 2006 through December 2021 were used for training and 47 observations from January 2022 through November 2025 were held out for final interim testing. The derived commodity-region-month table has 30,240 rows, 15 regions, eight commodity groups, no duplicate keys, and full climate coverage. It is an integration scaffold; national price variables repeat across regions and are labeled as proxies.",
    )
    add_caption(doc, "Table 5", "Core data dictionary (modeling variables)")
    add_table(
        doc,
        ["Variable", "Definition", "Type/unit", "Missing-value rule", "Role"],
        [
            ["date", "First day of observation month", "Date", "Not allowed", "Index"],
            ["region", "State/representative regional point", "Category", "Not allowed", "Panel key"],
            ["commodity", "One of eight essential-food groups", "Category", "Not allowed", "Panel key"],
            ["food_cpi_2015_100", "India food CPI, 2015=100", "Float/index", "Not imputed", "Target level"],
            ["food_cpi_mom_pct", "Monthly CPI percentage change", "Float/%", "First row undefined", "Shock input"],
            ["world_rice_usd_per_mt", "World Bank Thai 5% rice benchmark", "Float/USD per mt", "Time interpolation only if isolated", "Predictor"],
            ["world_wheat_usd_per_mt", "World Bank US SRW wheat benchmark", "Float/USD per mt", "Linear interpolation; documented 2.78%", "Predictor"],
            ["rainfall_mm", "NASA POWER precipitation parameter", "Float/mm per day", "-999 converted to missing", "Predictor"],
            ["temperature_c", "NASA POWER 2m temperature", "Float/Celsius", "-999 converted to missing", "Predictor"],
            ["relative_humidity_pct", "NASA POWER 2m relative humidity", "Float/%", "-999 converted to missing", "Predictor"],
            ["producer_price_lcu_per_tonne", "Median matched FAOSTAT annual producer price", "Float/INR per tonne", "Retain missing; no fabricated price", "Panel proxy"],
            ["production_tonnes", "Matched FAOSTAT annual production", "Float/tonnes", "Retain missing; audit aggregation", "Panel proxy"],
        ],
        widths=[1.35, 2.0, 1.0, 1.35, 0.8],
        font_size=7.5,
    )
    add_heading(doc, "GitHub Data Availability Statement", 2)
    add_bullets(
        doc,
        [
            "Raw acquisition scripts: scripts/01_extract_mospi_cpi_cfpi.py through scripts/09_build_interim_research_dataset.py.",
            "Curated, evaluator-accessible data: data/curated/*.csv.gz.",
            "Quality evidence and checksums: data/metadata/interim_data_manifest.csv and interim_data_quality_report.json.",
            "Preliminary analysis outputs: reports/tables/*.csv and reports/figures/*.png.",
            "Large raw archives: reproducibly downloaded to data/raw/ and excluded from Git history under the documented storage policy.",
        ],
    )

    add_heading(doc, "Analysis", 1)
    add_heading(doc, "Data Cleaning", 2)
    add_paragraph(
        doc,
        "Cleaning was conservative and source-aware. Area was filtered exactly to India before export from each FAOSTAT archive. Numeric values were parsed without replacing official flags. NASA annual summary keys (month code 13) were excluded from the monthly panel, and the API sentinel -999 was converted to missing. World Bank column labels were stripped of trailing spaces; date keys such as 1960M01 were converted to month starts. Duplicate checks were applied to the national row and the date-region-commodity key. Commodity mappings used documented regular-expression groups and remain reviewable in code.",
    )
    add_caption(doc, "Table 6", "Data cleaning log")
    add_table(
        doc,
        ["Issue", "Affected variables", "Detection", "Treatment", "Rationale"],
        [
            ["Nonmonthly NASA summaries", "All climate fields", "Month code 13", "Excluded", "Prevents invalid dates and double counting."],
            ["Missing World Bank wheat values", "Wheat benchmark", "2.78% audit", "Time interpolation in model copy", "Preserves a continuous exogenous series; raw values unchanged."],
            ["Lag/target edge missingness", "CPI lags and next target", "Null audit", "Rows dropped only for modeling", "Mechanically undefined observations cannot be imputed."],
            ["Duplicate keys", "Date-region-commodity", "Exact duplicate test", "None found", "Confirms one row per intended panel unit."],
            ["Source sentinels", "NASA parameters", "Value equals -999", "Converted to missing", "Sentinel is not a physical measurement."],
            ["National proxy repetition", "CPI/benchmark variables", "Lineage review", "Retained with explicit label", "Useful for prototype; prevents false regional claims."],
        ],
        widths=[1.2, 1.25, 1.15, 1.2, 1.7],
        font_size=8,
    )
    add_heading(doc, "Exploratory Data Analysis", 2)
    add_figure(doc, "food_cpi_trend.png", "Figure 1", "India food CPI trend, 2005-2025")
    add_paragraph(
        doc,
        "As shown in Figure 1, the food CPI displays a strong long-run increase with short periods of stabilization and renewed acceleration. This level behavior explains why a model that cannot extrapolate beyond its training range is structurally vulnerable. The figure also supports the decision to predict one month ahead using current levels and lagged values while evaluating errors strictly in future periods.",
    )
    add_figure(doc, "regional_rainfall_heatmap.png", "Figure 2", "Regional rainfall variation across the 15 selected points")
    add_paragraph(
        doc,
        "Figure 2 demonstrates substantial regional and temporal variation in the NASA POWER precipitation parameter. The variation is valuable for the planned market panel, but the interim national CPI model averages the 15 points. Therefore, the absence of a strong national rainfall coefficient does not imply that rainfall is unimportant for a particular crop or market. Aggregation can conceal localized and lagged effects.",
    )
    add_caption(doc, "Table 7", "EDA insight summary")
    add_table(
        doc,
        ["Reference", "What it shows", "Key insight", "Decision/next step"],
        [
            ["Figure 1", "Monthly food CPI level", "Strong trend and persistence", "Retain persistence control; use differenced targets and trend-capable models."],
            ["Figure 2", "Regional rainfall by year", "Climate exposure is heterogeneous", "Use crop-region lags after AGMARKNET backfill."],
            ["Table 1", "Data volume and coverage", "Large raw sources yield a compact auditable India subset", "Keep raw locally; commit curated evaluator data."],
            ["Quality audit", "Missingness and key integrity", "No duplicate panel keys; limited benchmark missingness", "Document interpolation and preserve raw values."],
        ],
        widths=[1.0, 1.6, 1.9, 2.0],
        font_size=8,
    )

    add_heading(doc, "Research Design and Methodology", 1)
    add_heading(doc, "Design", 2)
    add_paragraph(
        doc,
        "The study uses a quantitative, observational, longitudinal research design. Predictors at month t are used to estimate price level or shock status at t+1, with planned extensions to t+2 and t+3. No random assignment or causal treatment is present. The experimental comparison concerns algorithms evaluated on identical time periods and feature information. Persistence is the control model; ridge regression and random forest are interim experimental models. Later candidates include seasonal ARIMA, gradient boosting, XGBoost or LightGBM, and sequence models only if the market panel becomes sufficiently dense.",
    )
    add_heading(doc, "End-to-End Framework", 2)
    add_bullets(
        doc,
        [
            "Acquire official files and API responses; retain immutable raw data and record checksums.",
            "Filter and harmonize country, month, commodity, region, units, and source flags.",
            "Audit missingness, duplicates, outliers, structural breaks, and proxy limitations.",
            "Engineer lags, rolling summaries, seasonal sine/cosine terms, climate variables, production variables, and price benchmarks.",
            "Train models on past observations only; select hyperparameters inside time-series folds.",
            "Evaluate future holdouts with regression, classification, calibration, and paired-loss tests.",
            "Explain predictions, construct HFASI, and backtest decisions under alternative costs.",
        ],
    )
    add_heading(doc, "Features and Model Choice", 2)
    add_caption(doc, "Table 8", "Feature set and usage")
    add_table(
        doc,
        ["Feature group", "Examples", "Reason for inclusion", "Models"],
        [
            ["Autoregressive", "CPI level; lags 1, 3, 12", "Captures persistence, short memory, and annual seasonality.", "All regression/classification models"],
            ["International prices", "Rice, wheat, palm oil, soybean oil, sugar", "Represents external commodity pressure and common shocks.", "Ridge, random forest, OLS"],
            ["Energy", "Crude oil benchmark", "Proxy for transport, fertilizer, and production costs.", "Ridge, random forest"],
            ["Climate", "Rainfall, temperature, humidity", "Represents crop and supply exposure.", "Ridge, random forest, OLS"],
            ["Seasonality", "Month sine and cosine", "Cyclic encoding avoids arbitrary month ordering.", "Ridge, random forest, OLS"],
            ["Annual agriculture", "Producer price and production", "Links crop supply and farm-gate conditions.", "Final commodity-region panel"],
        ],
        widths=[1.2, 1.5, 2.4, 1.3],
        font_size=8,
    )
    add_paragraph(
        doc,
        "Ridge regression is appropriate because the predictors are correlated, the sample is modest, and coefficients remain more stable than unregularized estimates. Random forest tests nonlinear interactions and provides impurity-based feature importance, but it cannot extrapolate a rising level outside the training range. The interim failure of the random forest is therefore informative rather than a reason to hide the model. It directs the final study toward differenced targets, boosting models with trend features, and explicit statistical baselines.",
    )
    add_heading(doc, "Validation and Statistical Methods", 2)
    add_paragraph(
        doc,
        "The primary holdout begins in January 2022. All model fitting uses observations dated before that point, leaving 47 one-month-ahead test cases through November 2025. Five expanding time-series folds estimate training-period variability. No random shuffling is used. The paired absolute errors of the best model and persistence are compared with a one-sided Wilcoxon signed-rank test. For RQ1, standardized predictors are entered into an ordinary least squares model with heteroskedasticity-and-autocorrelation-consistent standard errors using three lags. Statistical significance is reported with 95% confidence intervals, but coefficients are interpreted as associations.",
    )
    add_caption(doc, "Table 9", "Evaluation metrics and formulae")
    add_table(
        doc,
        ["Metric", "Formula", "Interpretation", "Use"],
        [
            ["MAE", "(1/n) sum |y_i - yhat_i|", "Average absolute error in index points.", "Regression"],
            ["RMSE", "sqrt[(1/n) sum (y_i - yhat_i)^2]", "Penalizes large forecast errors.", "Regression"],
            ["sMAPE", "(100/n) sum 2|y-yhat|/(|y|+|yhat|)", "Scale-free symmetric percentage error.", "Regression"],
            ["R-squared", "1 - SSE/SST", "Share of holdout variance explained; can be negative.", "Regression"],
            ["Precision", "TP/(TP+FP)", "Share of warnings that are true shocks.", "Classification"],
            ["Recall", "TP/(TP+FN)", "Share of true shocks detected.", "Classification"],
            ["F1", "2PR/(P+R)", "Balance of precision and recall.", "Classification"],
            ["ROC-AUC", "P(score_positive > score_negative)", "Ranking discrimination across thresholds.", "Classification"],
        ],
        widths=[1.0, 2.2, 2.2, 1.0],
        font_size=8,
    )

    add_heading(doc, "Preliminary Results", 1)
    add_heading(doc, "Forecast Performance", 2)
    add_caption(doc, "Table 10", "One-month-ahead holdout performance, 2022-2025")
    add_table(
        doc,
        ["Model", "MAE", "RMSE", "sMAPE", "R-squared", "Key takeaway"],
        [
            ["Ridge regression", "1.362", "1.933", "0.910%", "0.957", "Best point estimates; small gain over persistence."],
            ["Persistence", "1.426", "2.045", "0.952%", "0.952", "Strong control because food CPI is highly persistent."],
            ["Random forest", "19.750", "21.948", "14.011%", "-4.487", "Failed to extrapolate to higher post-2021 CPI levels."],
        ],
        widths=[1.25, 0.65, 0.65, 0.75, 0.75, 2.15],
        font_size=8,
    )
    add_figure(doc, "model_performance_rmse.png", "Figure 3", "Comparison of holdout RMSE")
    add_paragraph(
        doc,
        "Table 10 and Figure 3 show that ridge regression produced the lowest holdout error, but the practical improvement over persistence was modest: RMSE declined by 0.112 index points. The one-sided paired Wilcoxon test did not establish a significant reduction in absolute error (p = .202). Five-fold rolling-origin MAE averaged 1.657 for ridge regression and 9.900 for random forest. The evidence therefore does not yet support rejecting the RQ2 null hypothesis that machine-learning approaches fail to improve materially on a conventional baseline.",
    )
    add_figure(doc, "actual_vs_predicted_food_cpi.png", "Figure 4", "Actual and ridge-predicted food CPI in the holdout period")
    add_paragraph(
        doc,
        "Figure 4 indicates that ridge predictions follow the rising level closely but smooth some monthly movements. The model is more useful as a transparent benchmark than as a finished early-warning tool. Prediction intervals, multi-horizon errors, and commodity-level tests remain necessary before operational use.",
    )
    add_heading(doc, "Driver Analysis", 2)
    add_figure(doc, "random_forest_feature_importance.png", "Figure 5", "Random forest feature importance")
    add_paragraph(
        doc,
        "Figure 5 shows that the random forest relied almost entirely on the current CPI level and its lags. This concentration explains both its historical fit and its poor extrapolation. In the HAC-robust OLS of next-month percentage change (n = 250; adjusted R-squared = .341), standardized world wheat price (b = 0.256, p = .005), temperature (b = 0.777, p = .006), and world rice price (b = 0.127, p = .049) were positively associated with the next monthly change. The lagged CPI coefficient was negative (b = -0.152, p = .008), consistent with short-run mean reversion. Rainfall was not significant at the national aggregation level (p = .254). These estimates are preliminary associations and may change after region-commodity disaggregation.",
    )
    add_heading(doc, "Shock Classification", 2)
    add_paragraph(
        doc,
        "The hybrid shock label identifies an event when the next monthly change is at least two rolling standard deviations from its recent mean or exceeds the approved 10% economic threshold. Only three positive events occurred in the 47-row test set. At the default 0.50 probability threshold, the random forest classifier achieved accuracy .851 and ROC-AUC .818, but precision, recall, and F1 were all zero because no positive event crossed the decision threshold. Accuracy is therefore misleading. The next iteration will tune thresholds inside training folds, emphasize precision-recall curves, compare class-weighted logistic regression and gradient boosting, and evaluate warning cost rather than default classification alone.",
    )
    add_heading(doc, "Findings by Research Question", 2)
    add_caption(doc, "Table 11", "Interim findings and remaining evidence")
    add_table(
        doc,
        ["RQ", "What is supported now", "What remains"],
        [
            ["RQ1", "Lagged CPI and selected global price/climate variables show predictive associations in the national prototype.", "Regional market prices, crop-specific lags, multicollinearity audit, stability and causal caution."],
            ["RQ2", "Ridge slightly beats persistence; random forest fails under level extrapolation.", "No significant superiority yet; broaden models, horizons, intervals, and paired tests."],
            ["RQ3", "Data architecture and HCES food-share weights are defined.", "Construct, calibrate, validate, and stress-test HFASI by segment."],
            ["RQ4", "Decision users and evaluation losses are specified.", "Backtest household, retail, procurement, and policy decisions against baselines."],
        ],
        widths=[0.55, 3.0, 2.95],
        font_size=8,
    )

    add_heading(doc, "Interim Limitations and Risks", 1)
    add_bullets(
        doc,
        [
            "The analytical target is national food CPI, not a completed mandi-price panel. National variables repeated across regions are proxies and must not be interpreted as local market prices.",
            "The AGMARKNET API requires an authorized data.gov.in key. The collector is present, but historical records were not fabricated or accessed with an unauthorized credential.",
            "The World Bank workbook available through the cited page was updated January 2025; 2025 values therefore have limited end-period coverage for wheat and are documented.",
            "Rare shocks make threshold-specific precision and recall unstable. Three holdout events are insufficient for a final operational claim.",
            "The random forest cannot extrapolate rising price levels and should be evaluated on changes or with trend-capable models.",
            "Climate points represent selected regional locations, not complete state-area averages. Sensitivity to spatial aggregation is pending.",
            "Statistical associations may reflect common trends, omitted variables, policy interventions, or reverse timing; no causal effect is claimed.",
            "FAOSTAT and NASA are authoritative standardized sources but do not replace India-specific official market and household microdata in the final analysis.",
        ],
    )

    add_heading(doc, "Next Steps for the Final Report", 1)
    add_bullets(
        doc,
        [
            "Obtain and store the user-authorized data.gov.in API key in .env; backfill AGMARKNET prices and arrivals without committing credentials.",
            "Harmonize market, district, state, commodity, unit, and month keys; publish coverage and attrition tables.",
            "Extract HCES rural, urban, and expenditure-fractile food shares; implement HFASI with baseline 100 and sensitivity bands.",
            "Model price changes and levels at one-, two-, and three-month horizons using persistence, seasonal naive, ARIMA/ETS, ridge, gradient boosting, and selected ensembles.",
            "Use nested rolling-origin validation, prediction intervals, Diebold-Mariano or paired bootstrap comparisons, and calibration diagnostics.",
            "Tune shock thresholds inside training data; report precision-recall AUC, recall at fixed alert budgets, Brier score, and expected decision cost.",
            "Generate SHAP explanations and compare their stability across time, commodities, and regions.",
            "Backtest procurement, inventory, household-budget, and policy-monitoring scenarios; document assumptions and failure modes.",
            "Complete model cards, ethical review, accessibility checks, tests, and a reproducible final release.",
        ],
    )

    add_heading(doc, "Bibliography", 1)
    references = [
        "Akter, S., & Basher, S. A. (2014). The impacts of food price and income shocks on household food security and economic well-being: Evidence from rural Bangladesh. Global Environmental Change, 25, 150-162. https://doi.org/10.1016/j.gloenvcha.2014.02.003",
        "Bhardwaj, M. R., Pawar, J., Bhat, A., Deepanshu, Enaganti, I., Sagar, K., & Narahari, Y. (2023). An innovative deep learning based approach for accurate agricultural crop price prediction. 2023 IEEE 19th International Conference on Automation Science and Engineering. https://doi.org/10.1109/CASE56687.2023.10260494",
        "Bogmans, C., Pescatori, A., & Prifti, E. (2024). How do economic growth and food inflation affect food insecurity? IMF Working Paper No. 2024/188. https://doi.org/10.5089/9798400287336.001",
        "Cattaneo, A., Sadiddin, A., Vaz, S., Conti, V., Holleman, C., Sanchez, M. V., & Torero, M. (2023). Ensuring affordability of diets in the face of shocks. Food Policy, 117, 102470. https://doi.org/10.1016/j.foodpol.2023.102470",
        "Food and Agriculture Organization of the United Nations. (2026). FAOSTAT bulk data: Consumer price indices, producer prices, and crop and livestock production. https://www.fao.org/faostat/en/",
        "Hyndman, R. J., & Koehler, A. B. (2006). Another look at measures of forecast accuracy. International Journal of Forecasting, 22(4), 679-688. https://doi.org/10.1016/j.ijforecast.2006.03.001",
        "Jain, A., Marvaniya, S., Godbole, S., & Munigala, V. (2020). A framework for crop price forecasting in emerging economies by analyzing the quality of time-series data. arXiv. https://doi.org/10.48550/arXiv.2009.04171",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765-4774.",
        "Ma, W., Nowocin, K., Marathe, N., & Chen, G. H. (2019). An interpretable produce price forecasting system for small and marginal farmers in India using collaborative filtering and adaptive nearest neighbors. Proceedings of ICTD, Article 6. https://doi.org/10.1145/3287098.3287100",
        "Madaan, L., Sharma, A., Khandelwal, P., Goel, S., Singla, P., & Seth, A. (2019). Price forecasting and anomaly detection for agricultural commodities in India. Proceedings of COMPASS. https://doi.org/10.1145/3314344.3332488",
        "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2018). Statistical and machine learning forecasting methods: Concerns and ways forward. PLOS ONE, 13(3), e0194889. https://doi.org/10.1371/journal.pone.0194889",
        "Malesios, C., Jones, N., & Jones, A. (2020). A change-point analysis of food price shocks. Climate Risk Management, 27, 100208. https://doi.org/10.1016/j.crm.2019.100208",
        "Ministry of Agriculture and Farmers Welfare, Directorate of Economics and Statistics. (2025). Agricultural prices in India 2024. https://desagri.gov.in/document-report/agricultural-prices-in-india-2024/",
        "Ministry of Statistics and Programme Implementation. (2024). Survey on Household Consumption Expenditure: 2022-23. Government of India.",
        "NASA POWER Project. (2026). Monthly and annual API. https://power.larc.nasa.gov/docs/services/api/temporal/monthly/",
        "Theofilou, A., Arvanitakis, G., & Tzouramani, I. (2025). Predicting prices of staple crops using machine learning: A systematic review of studies on wheat, corn, and rice. Sustainability, 17(12), 5456. https://doi.org/10.3390/su17125456",
        "World Bank. (2025). World Bank commodity price data (The Pink Sheet): Historical monthly data. https://www.worldbank.org/en/research/commodity-markets",
    ]
    for reference in references:
        p = add_paragraph(doc, reference, no_indent=True)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    add_heading(doc, "Appendix A: Reproducibility and Evidence Map", 1)
    add_caption(doc, "Table A1", "Repository evidence map")
    add_table(
        doc,
        ["Evidence", "Path"],
        [
            ["Source checksums and row counts", "data/metadata/interim_data_manifest.csv"],
            ["Missingness and key audit", "data/metadata/interim_data_quality_report.json"],
            ["Model summary and statistical comparison", "data/metadata/interim_analysis_summary.json"],
            ["Forecast performance table", "reports/tables/preliminary_regression_performance.csv"],
            ["Shock classification table", "reports/tables/preliminary_shock_classification_performance.csv"],
            ["HAC-robust coefficients", "reports/tables/preliminary_ols_hac_coefficients.csv"],
            ["Acquisition and data build", "scripts/09_build_interim_research_dataset.py"],
            ["Analysis and figures", "scripts/10_run_interim_analysis.py"],
        ],
        widths=[2.6, 3.9],
        font_size=9,
    )
    add_heading(doc, "Appendix B: HFASI Planned Formula", 1)
    add_paragraph(
        doc,
        "For household segment s and month t, the planned unscaled affordability pressure is AP_s,t = w_food,s x (PredictedFoodCostGrowth_t - ExpectedPurchasingPowerGrowth_s,t), where w_food,s is the segment food-expenditure share. HFASI will map this pressure to a baseline-100 index with transparent caps and sensitivity analysis. The final report will test internal consistency, rank stability, criterion validity against observed CPI/CFPI and expenditure outcomes, and alternative rural/urban weights. No interim HFASI score is reported because calibration without completed HCES extraction would create false precision.",
    )

    # Ensure Word refreshes PAGE fields on open.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
