"""Build the final QM640 interim report from the retained Walsh template."""

from __future__ import annotations

import copy
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
TEMPLATE = Path(r"C:\Users\piyus\Downloads\QM 640 Interim Report template-1 (2).docx")
OUTPUT = ROOT / "reports" / "Piyush_Soni_QM640_Interim_Report_Final.docx"
ASSETS = WORKSPACE / "tmp" / "interim_report" / "assets"
SKILL_SCRIPTS = Path(
    r"C:\Users\piyus\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.727.11326\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "E7E6E6"
MID_GRAY = "BFBFBF"
REPO = "https://github.com/piyushsoni88/qm640-food-affordability-ai"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=12, bold=None, italic=None, color=BLACK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color


def format_paragraph(
    paragraph,
    *,
    first_indent=True,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    double=True,
    before=0,
    after=0,
    keep_next=False,
) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = align
    pf.first_line_indent = Inches(0.5) if first_indent else None
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = keep_next
    pf.widow_control = True
    if double:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing = 1.0
    for run in paragraph.runs:
        set_run_font(run)


def add_body(doc, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    format_paragraph(p)
    return p


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, bold=True)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = None
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0 if level == 1 else 2)
    pf.space_after = Pt(0)
    pf.keep_with_next = True
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def add_hyperlink(paragraph, url: str, text: str | None = None):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.extend([fonts, color, underline, size])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text or url
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, size=10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def clear_container(container) -> None:
    element = container._element
    for child in list(element):
        element.remove(child)


def configure_headers(doc) -> None:
    for idx, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        clear_container(section.header)
        clear_container(section.footer)
        table = section.header.add_table(rows=1, cols=2, width=Inches(6.5))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        widths = [4680, 4680]
        apply_table_geometry(
            table, widths, table_width_dxa=9360, indent_dxa=0,
            cell_margins_dxa={"top": 0, "bottom": 0, "start": 0, "end": 0},
        )
        left, right = table.rows[0].cells
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if idx > 0:
            run = lp.add_run("Interim Report")
            set_run_font(run, size=10)
        add_page_field(rp)
        for p in [lp, rp]:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0


def add_table_caption(doc, number: int, title: str):
    p = doc.add_paragraph()
    r1 = p.add_run(f"Table {number}")
    set_run_font(r1, bold=True)
    r1.add_break()
    r2 = p.add_run(title)
    set_run_font(r2, italic=True)
    format_paragraph(
        p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT,
        double=False, after=3, keep_next=True,
    )
    return p


def add_figure_caption(doc, number: int, title: str):
    p = doc.add_paragraph()
    r1 = p.add_run(f"Figure {number}")
    set_run_font(r1, bold=True)
    r1.add_break()
    r2 = p.add_run(title)
    set_run_font(r2, italic=True)
    format_paragraph(
        p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT,
        double=False, after=3, keep_next=True,
    )
    return p


def add_note(doc, text: str):
    p = doc.add_paragraph()
    r1 = p.add_run("Note. ")
    set_run_font(r1, size=9, italic=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=9)
    format_paragraph(
        p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT,
        double=False, before=3, after=4,
    )
    return p


def add_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    weights: list[float],
    *,
    font_size=9,
    header_fill=LIGHT_GRAY,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(value)
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_run_font(run, size=font_size, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[idx].paragraphs:
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if len(str(value)) < 18 and idx > 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=font_size)
    widths = column_widths_from_weights(weights, 9360)
    apply_table_geometry(
        table, widths, table_width_dxa=9360, indent_dxa=0,
        cell_margins_dxa={"top": 70, "bottom": 70, "start": 90, "end": 90},
    )
    return table


def add_picture(doc, path: Path, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    format_paragraph(
        p, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER,
        double=False, before=2, after=2, keep_next=True,
    )
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_reference(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return p


def add_repo_link(doc, label: str, path: str, explanation: str):
    p = doc.add_paragraph()
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True)
    add_hyperlink(p, f"{REPO}/{path}", f"{REPO}/{path}")
    tail = p.add_run(f" - {explanation}")
    set_run_font(tail)
    format_paragraph(p, first_indent=False)
    return p


def prepare_document() -> Document:
    doc = Document(TEMPLATE)
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    # The retained final section becomes the title-page section.
    first = doc.sections[0]
    first.page_width = Inches(8.5)
    first.page_height = Inches(11)
    first.top_margin = Inches(1)
    first.bottom_margin = Inches(1)
    first.left_margin = Inches(1)
    first.right_margin = Inches(1)
    first.header_distance = Inches(0.5)
    first.footer_distance = Inches(0.5)

    # Apply template-derived APA styles.
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.first_line_indent = Inches(0.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for level in [1, 2, 3]:
        style_name = f"Heading {level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True
    try:
        doc.styles["List Bullet"]
    except KeyError:
        bullet = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet.base_style = normal
        bullet.paragraph_format.left_indent = Inches(0.5)
        bullet.paragraph_format.first_line_indent = Inches(-0.25)
    try:
        doc.styles["Table Grid"]
    except KeyError:
        doc.styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)
    return doc


def add_title_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    items = [
        ("Data Analytics Capstone", True, 14),
        ("Forecasting Essential Food Price Shocks and Household Affordability Stress in India", True, 14),
        ("An Explainable AI Decision-Intelligence Framework Using Public Market, Agricultural, Climatic, and Economic Data", False, 13),
        ("Interim Report", True, 14),
        ("Piyush Soni", False, 12),
        ("Walsh College", False, 12),
        ("QM640: Data Analytics Capstone", False, 12),
        ("Mentor: Prof. Rishabh Pandey", False, 12),
        ("Summer 2026 Term", False, 12),
        ("July 28, 2026", False, 12),
    ]
    for idx, (text, bold, size) in enumerate(items):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = None
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(7 if idx in {3, 4} else 0)
        pf.space_after = Pt(0)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.page_width = Inches(8.5)
    body_section.page_height = Inches(11)
    body_section.top_margin = Inches(1.5)
    body_section.bottom_margin = Inches(1.5)
    body_section.left_margin = Inches(1)
    body_section.right_margin = Inches(1)
    body_section.header_distance = Inches(0.5)
    body_section.footer_distance = Inches(0.5)


def build_report() -> Document:
    doc = prepare_document()
    add_title_page(doc)
    configure_headers(doc)

    # Page 2 - GitHub and status
    add_heading(doc, "GitHub Repository and Project Status", 1)
    p = doc.add_paragraph()
    p.add_run(
        "The public repository containing the permitted data, acquisition and "
        "analysis code, notebooks, documentation, and report materials is: "
    )
    for run in p.runs:
        set_run_font(run)
    add_hyperlink(p, REPO, REPO)
    format_paragraph(p)
    add_repo_link(doc, "Curated data", "tree/main/data/curated", "compressed evaluator-accessible analytical source files")
    add_repo_link(doc, "Data dictionary", "blob/main/data/data_dictionary.csv", "variable definitions and transformations")
    add_repo_link(doc, "Source manifest", "blob/main/data/source_manifest.csv", "source URLs, provenance, and coverage")
    add_repo_link(doc, "Executed workflow", "tree/main/notebooks", "ten Colab-compatible notebooks from acquisition through consolidation")
    add_repo_link(doc, "Reusable scripts", "tree/main/scripts", "parameterized collection, transformation, and report builders")
    add_repo_link(doc, "Report outputs", "tree/main/reports", "figures, tables, model cards, and Word report")
    add_heading(doc, "Interim Project Status (Progress Snapshot)", 2)
    add_bullet(doc, "Completed: public-data acquisition, lineage controls, cleaning, exploratory analysis, statistical modeling, forecasting, shock classification, explainability, affordability indexing, stress scenarios, and final evidence consolidation.")
    add_bullet(doc, "Completed: 10 sequential notebooks executed successfully; Notebook 10 verified nine upstream summaries and inventoried 98 report artifacts.")
    add_bullet(doc, "In progress: final narrative packaging and synchronized GitHub publication of the executed notebooks and report.")
    add_bullet(doc, "Pending for the final study: licensed HCES microdata extraction, independent household-level validation, broader decision-loss backtesting, and monitored deployment testing.")
    add_table_caption(doc, 1, "Interim analytical snapshot")
    add_table(
        doc,
        ["Evidence layer", "Verified result", "Coverage / interpretation"],
        [
            ["Curated analytical rows", "1,494,745", "Permitted compressed files across official and multilateral sources"],
            ["AGMARKNET aggregate", "969,377 rows", "Represents 18,836,462 underlying official price observations"],
            ["Clean state panel", "41,792 rows", "32 reporting states/UTs; 8 commodities; 2001-2026 YTD"],
            ["National panel", "319 months", "January 2000-July 2026; food CPI observed through December 2025"],
        ],
        [2.1, 1.4, 3.0],
        font_size=9,
    )
    add_note(doc, "Large or restricted raw archives remain local and are recreated through documented scripts; no credentials or unauthorized microdata are committed.")
    add_page_break(doc)

    # Page 3 - Introduction
    add_heading(doc, "Introduction", 1)
    add_heading(doc, "Background and Context", 2)
    add_body(
        doc,
        "Food-price volatility is simultaneously a household-welfare problem, a procurement risk, and a policy-monitoring challenge. Rice, wheat, onion, potato, tomato, pulses, edible oil, and sugar differ in production cycles, storability, market concentration, and transport exposure. Their prices can respond to crop losses, market-arrival contractions, fuel costs, trade restrictions, festival demand, and storage constraints. When signals are detected only after retail prices rise, households lose purchasing power and food-sector enterprises face margin compression, stock-outs, or emergency purchases."
    )
    add_body(
        doc,
        "The welfare effect is unequal because food absorbs a larger share of household expenditure in rural and lower-income settings. The Household Consumption Expenditure Survey (HCES) 2022-2023 reports food shares of approximately 46.3% of rural Monthly Per Capita Consumption Expenditure and 39.1% of urban expenditure (Ministry of Statistics and Programme Implementation [MOSPI], 2024). A useful system must therefore do more than forecast an index: it must disclose uncertainty, explain predictive signals, and translate expected cost pressure into a segment-sensitive affordability measure."
    )
    add_body(
        doc,
        "India offers rich but fragmented public data. AGMARKNET provides official mandi observations; FAOSTAT supplies standardized agricultural and price series; NASA POWER provides reproducible climate histories; and the World Bank Pink Sheet provides monthly global commodity and energy benchmarks. Combining these sources supports a longitudinal state-commodity panel and a national forecasting panel, while also introducing differences in frequency, units, reporting completeness, and release timing that require explicit lineage and quality controls."
    )
    add_heading(doc, "Problem Statement", 2)
    add_body(
        doc,
        "Indian households and food-sector enterprises face recurring uncertainty because essential-food prices are volatile, geographically uneven, and shaped by interacting climatic, seasonal, agricultural, market, and macroeconomic forces. Although relevant public data exist, they are commonly used for retrospective monitoring in separate systems. The operational gap is an integrated and explainable early-warning workflow that compares transparent baselines with machine learning, quantifies uncertainty, detects abnormal inflation risk, and converts forecast food-cost pressure into an interpretable affordability measure."
    )
    add_heading(doc, "Purpose of the Study", 2)
    add_body(
        doc,
        "This quantitative longitudinal study develops and validates a reproducible decision-intelligence framework for essential-food price risk and household affordability stress in India. It connects four tasks: identify conditional drivers, forecast national food CPI, classify next-month high-inflation shocks, and estimate representative rural and urban affordability stress. Chronological validation prevents look-ahead bias; permutation importance and partial dependence support predictive interpretation; and scenario analysis demonstrates how household-budget and procurement signals change under controlled assumptions."
    )
    add_page_break(doc)

    # Page 4 - Scope and research questions
    add_heading(doc, "Scope and Objectives", 1)
    add_body(
        doc,
        "The implemented scope covers eight essential commodities and all available reporting states and union territories in the official market-price panel. The cleaned state-month-commodity panel contains 41,792 observations from January 2001 through July 2026. A national monthly panel contains 319 rows from January 2000 through July 2026, with the food CPI target observed through December 2025. Analyses compare recent two-, five-, and ten-year regional windows; one-step and 24-step forecasts; next-month shock classification; and January-December 2026 affordability and stress scenarios."
    )
    add_heading(doc, "Research Question 1 (RQ1)", 2)
    add_body(doc, "Which lagged-price, market, climatic, seasonal, and macroeconomic variables are conditionally associated with food-price change and high-inflation risk?")
    add_heading(doc, "Research Question 2 (RQ2)", 2)
    add_body(doc, "Do machine-learning and ensemble models predict national food CPI more accurately than transparent statistical and naive baselines under chronological validation?")
    add_heading(doc, "Research Question 3 (RQ3)", 2)
    add_body(doc, "How can forecast food-cost growth be translated into a transparent Household Food Affordability Stress Index (HFASI) for representative rural and urban segments?")
    add_heading(doc, "Research Question 4 (RQ4)", 2)
    add_body(doc, "How can combined shock probabilities and affordability stress support household-budget, procurement, inventory, and policy-monitoring decisions under alternative scenarios?")
    add_heading(doc, "Measurable Objectives", 2)
    add_bullet(doc, "Create an auditable public-data pipeline with source manifests, data dictionary, quality flags, exact calendar lags, and reproducible Colab notebooks.")
    add_bullet(doc, "Compare persistence, seasonal-naive, regularized regression, tree ensembles, and damped exponential smoothing with expanding chronological validation.")
    add_bullet(doc, "Evaluate rare-event classification using balanced accuracy, precision, recall, specificity, F1, ROC-AUC, average precision, Brier score, and a confusion matrix.")
    add_bullet(doc, "Construct HFASI with published rural/urban food shares, explicit purchasing-power assumptions, forecast uncertainty, and sensitivity analysis.")
    add_bullet(doc, "Demonstrate an evidence-bounded decision framework without converting predictive associations into causal claims or automatic policy triggers.")
    add_page_break(doc)

    # Pages 5-6 - Literature
    add_heading(doc, "Literature Survey", 1)
    add_heading(doc, "Literature Review Approach", 2)
    add_body(
        doc,
        "Sources were selected from peer-reviewed journals, conference proceedings, systematic reviews, and official statistical publications using combinations of agricultural price forecasting, food inflation, price shock, explainable machine learning, forecast evaluation, household affordability, and India. Inclusion required direct relevance to at least one research question, a described method, and a traceable publication. Studies relying on random train-test splits were retained as domain context but not treated as temporal-validation benchmarks."
    )
    add_table_caption(doc, 2, "Literature relevance matrix")
    literature_rows = [
        ["Akter & Basher (2014)", "Household welfare", "Shock/welfare analysis", "Food and income shocks reduce welfare; supports RQ3 segment sensitivity."],
        ["Cattaneo et al. (2023)", "Diet affordability", "Policy synthesis", "Links price shocks to diet affordability and mitigation; supports RQ3/RQ4."],
        ["Hyndman & Koehler (2006)", "Forecast evaluation", "Metric analysis", "Clarifies scale-dependent and scaled errors; supports RQ2 metrics."],
        ["Makridakis et al. (2018)", "Forecast comparison", "Empirical review", "Warns that complex ML does not always beat statistical baselines; supports RQ2."],
        ["Jain et al. (2020)", "Crop-price forecasting", "Quality-aware framework", "Makes data quality part of forecasting design; supports RQ1/RQ2."],
        ["Madaan et al. (2019)", "Indian commodities", "Forecasting/anomaly detection", "Shows value of joint forecast and anomaly workflows; supports RQ2/RQ3."],
        ["Ma et al. (2019)", "Indian produce prices", "Interpretable collaborative filtering", "Demonstrates farmer-facing interpretable price predictions; supports RQ2/RQ4."],
        ["Bhardwaj et al. (2023)", "Agricultural forecasting", "Deep learning", "Shows nonlinear predictive capacity but reinforces sample-size needs; supports RQ2."],
        ["Lundberg & Lee (2017)", "Model interpretation", "SHAP theory", "Provides a framework for local/global predictive explanation; supports RQ1."],
        ["Theofilou et al. (2025)", "Staple-crop forecasting", "Systematic review", "Identifies gaps in external variables, interpretability, and validation; supports all RQs."],
    ]
    add_table(
        doc,
        ["Source", "Domain / context", "Method", "Key finding and project relevance"],
        literature_rows,
        [1.25, 1.25, 1.35, 2.65],
        font_size=8.3,
    )
    add_page_break(doc)
    add_heading(doc, "Thematic Synthesis", 2)
    add_body(
        doc,
        "First, the forecasting literature favors strong baselines and time-aware evaluation. Hyndman and Koehler (2006) show that error measures answer different questions, while Makridakis et al. (2018) caution that algorithmic complexity is not synonymous with accuracy. This evidence directly motivated separate one-step and 24-step backtests, expanding rather than shuffled validation, and the requirement that a complex model beat naive persistence before promotion."
    )
    add_body(
        doc,
        "Second, agricultural forecasting depends on data quality and context. Jain et al. (2020) emphasize quality-aware time-series design, and Madaan et al. (2019) connect price prediction with anomaly detection in India. Ma et al. (2019) demonstrate the value of interpretable outputs for marginal users, whereas Bhardwaj et al. (2023) illustrates the capacity and data demands of deep learning. The present project therefore retains missingness and coverage indicators, avoids price interpolation, and limits model complexity to what chronological evidence supports."
    )
    add_body(
        doc,
        "Third, price shocks become decision relevant only when translated into exposure. Akter and Basher (2014) and Cattaneo et al. (2023) show that food and income shocks affect welfare through household budget shares and purchasing power. This supports HFASI as a formative scenario index rather than a generic price score. It also supports separate rural and urban weights and requires purchasing-power assumptions to accompany every reported value."
    )
    add_body(
        doc,
        "Finally, explainability must be interpreted carefully. Lundberg and Lee (2017) formalize additive model explanations, but feature attribution describes the fitted model rather than causal intervention effects. The systematic review by Theofilou et al. (2025) similarly identifies validation and interpretability gaps. Accordingly, this study uses out-of-sample permutation importance, partial dependence, local sensitivity, and explicit causal-language restrictions."
    )
    add_heading(doc, "Literature-to-Design Decisions", 2)
    add_bullet(doc, "Use last-value and seasonal-naive forecasts as mandatory baselines.")
    add_bullet(doc, "Use expanding-window validation and horizon-specific error reporting.")
    add_bullet(doc, "Preserve raw outliers and reporting gaps; distinguish data quality from price movement.")
    add_bullet(doc, "Evaluate rare shocks with recall, average precision, calibration, and confusion counts.")
    add_bullet(doc, "Report affordability and scenario assumptions alongside point predictions.")
    add_page_break(doc)

    # Page 7 - Data sources
    add_heading(doc, "Data Description", 1)
    add_heading(doc, "Data Sources and Access", 2)
    add_body(
        doc,
        "The study integrates openly accessible government and multilateral data. The largest source is the AGMARKNET historical market-price layer, stored as a compressed daily-state aggregate representing 18,836,462 underlying official observations. Climate histories come from NASA POWER; international price and energy series come from the World Bank Pink Sheet; standardized agricultural production and price indicators come from FAOSTAT; and HCES supplies published rural and urban food-expenditure shares. Each file is identified in the source manifest with URL, coverage, extraction metadata, and derived-file path."
    )
    add_table_caption(doc, 3, "Principal data sources and analytical roles")
    add_table(
        doc,
        ["Source", "Variables / coverage", "Analytical role", "Access"],
        [
            ["AGMARKNET / OGD India", "Daily min, max, modal price; market and commodity; 2000-2026 YTD", "State-commodity wholesale-price backbone", "data.gov.in / AGMARKNET"],
            ["NASA POWER", "Daily and monthly rainfall, temperature, humidity; all states/UTs", "Climate level and anomaly features", "power.larc.nasa.gov"],
            ["World Bank Pink Sheet", "Monthly wheat, rice, oils, sugar, crude oil", "International commodity and energy controls", "worldbank.org/commodities"],
            ["FAOSTAT", "Food CPI, producer prices, crop production", "Benchmark and annual supply indicators", "fao.org/faostat"],
            ["MOSPI HCES 2022-23", "Rural/urban expenditure and food shares", "HFASI representative segment weights", "mospi.gov.in"],
        ],
        [1.35, 2.4, 1.65, 1.1],
        font_size=8.5,
    )
    add_heading(doc, "Dataset Overview", 2)
    overview = [
        ["State analytical records", "41,792 cleaned state-month-commodity rows"],
        ["State/UT reporting coverage", "32 price-reporting regions; climate matching 100%"],
        ["Commodities", "8 essential commodities"],
        ["National records", "319 monthly rows; 312 with food CPI"],
        ["Time period", "State: Jan. 2001-Jul. 2026; national: Jan. 2000-Jul. 2026"],
        ["Units of analysis", "State-commodity-month and national month"],
        ["Targets", "Food CPI level; exact YoY/MoM change; next-month shock label; HFASI"],
    ]
    add_table(doc, ["Attribute", "Verified value"], overview, [2.0, 4.5], font_size=9)
    add_note(doc, "National food CPI is available through December 2025. January 2026 uses observed December 2025 predictors; later 2026 risks use conditional forecast inputs.")
    add_page_break(doc)

    # Page 8 - dictionary and availability
    add_heading(doc, "Data Dictionary (Mandatory)", 1)
    add_table_caption(doc, 4, "Core analytical data dictionary")
    add_table(
        doc,
        ["Variable", "Definition", "Type / unit", "Missing-value handling", "Role"],
        [
            ["date", "Monthly observation or forecast-origin date", "Date (YYYY-MM)", "Reject unparseable dates", "Panel key"],
            ["region / commodity", "Reporting state/UT and standardized food item", "Categorical", "Preserve explicit missing/unknown labels", "Keys / controls"],
            ["modal_price", "Representative official mandi price", "INR/quintal", "No interpolation; retain gaps", "Target / predictor"],
            ["food_cpi_2015_100", "National food CPI index", "Index, 2015=100", "Retain source missingness", "Forecast target"],
            ["rainfall / temperature anomaly", "Deviation from month-specific climate normal", "% / degrees C", "No zero-fill in observed data", "Predictors"],
            ["lag and rolling features", "Exact calendar lags and prior-window statistics", "Continuous", "Available only when keyed history exists", "Predictors"],
            ["shock_next", "Next-month YoY food CPI at/above training 75th percentile", "Binary", "Created only with observed target", "Classification target"],
            ["HFASI", "100 + food share x (food-cost growth - purchasing-power growth)", "Baseline-100 index", "Scenario assumptions required", "Decision output"],
        ],
        [1.2, 2.0, 1.0, 1.45, 0.85],
        font_size=8.2,
    )
    add_heading(doc, "GitHub Data Availability Statement", 2)
    add_body(
        doc,
        "Evaluator-accessible curated datasets are committed under data/curated, including the compressed AGMARKNET daily-state aggregate, state and national affordability panels, NASA climate panels, FAOSTAT extracts, and World Bank monthly indicators. The data dictionary and source manifest are stored at the root of data/. Raw bulk archives and credential-protected sources are excluded under the repository storage policy; retrieval scripts and checksums provide reproducibility without exposing credentials or violating source terms."
    )
    add_heading(doc, "Repository Evidence", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "The complete folder tree and notebook-to-output map appear in Appendix A. "
        "The repository README provides execution order, source policy, and environment instructions. "
    )
    set_run_font(r)
    add_hyperlink(p, f"{REPO}/blob/main/README.md", "Open README")
    format_paragraph(p)
    add_bullet(doc, "Raw data path: data/raw/ (local-only bulk or credential-dependent archives; retrieval code committed).")
    add_bullet(doc, "Curated/processed evaluator data: data/curated/.")
    add_bullet(doc, "Code and notebooks: scripts/ and notebooks/.")
    add_bullet(doc, "Figures, tables, and model cards: reports/.")
    add_page_break(doc)

    # Page 9 - Cleaning and EDA
    add_heading(doc, "Analysis", 1)
    add_heading(doc, "Data Cleaning", 2)
    add_body(
        doc,
        "Notebook 02 applied schema checks, date normalization, controlled region and commodity labels, numeric coercion, unit/range validation, duplicate-key audits, and exact calendar-lag construction. No price observation was interpolated. The validity rules removed zero rows from the 41,792-row state panel; 15 extreme-price rows (0.036%) were flagged while their raw values were preserved. Exact one-month lags were available for 40,189 rows and exact 12-month lags for 37,090 rows. Climate anomalies were calculated for 41,790 rows after deduplicating repeated state-month climate values across commodities."
    )
    add_table_caption(doc, 5, "Data cleaning log")
    add_table(
        doc,
        ["Issue", "Variables affected", "Detection", "Treatment", "Rationale"],
        [
            ["Missing prices", "Modal-price measures", "Null and coverage profiles", "Retained; no interpolation", "Avoid invented official prices"],
            ["Extreme prices", "State commodity prices", "Robust tail rules", "Flagged 15; raw values preserved", "Possible genuine shocks"],
            ["Duplicates", "State-commodity-month keys", "Exact key audit", "Deterministic aggregation/check", "Prevent duplicated weight"],
            ["Calendar gaps", "Lagged features", "Keyed date joins", "Lag missing unless exact month exists", "Prevent false row-offset lags"],
            ["Climate repetition", "State-month climate", "State-month deduplication", "One climate value per state-month", "Avoid commodity-count weighting"],
        ],
        [1.05, 1.35, 1.15, 1.5, 1.45],
        font_size=8.2,
    )
    add_heading(doc, "Exploratory Data Analysis Results", 2)
    add_body(
        doc,
        "The national series rises over the long run but contains distinct inflation surges and reversals. Onion has the highest commodity-level YoY volatility (standard deviation 117.75 percentage points), consistent with perishability and episodic supply disruption. Regional pressure rankings are stable in direction but differ in magnitude: West Bengal records the highest eligible median YoY pressure in the two-year (3.13%), five-year (6.70%), and ten-year (4.06%) windows. Raw rainfall-price and temperature-price correlations are near zero, indicating that contemporaneous bivariate relationships are insufficient for causal interpretation."
    )
    add_figure_caption(doc, 1, "National food CPI and official market-price index")
    add_picture(doc, ASSETS / "nb03_fig1.png", width=6.3)
    add_note(doc, "The national CPI trend is smoother than the volatile official mandi-price aggregate. Both series are indexed for comparison; the chart does not imply one causes the other.")
    add_page_break(doc)

    # Page 10 - EDA interpretation and modeling
    add_heading(doc, "EDA Interpretation and Modelling Design", 1)
    add_figure_caption(doc, 2, "Regional median price pressure across two-, five-, and ten-year windows")
    add_picture(doc, ASSETS / "nb03_fig4.png", width=6.35)
    add_note(doc, "Only regions meeting the coverage rule are ranked. The repeated West Bengal result is descriptive and may reflect commodity mix, reporting, market structure, or common shocks.")
    add_table_caption(doc, 6, "EDA insight summary")
    add_table(
        doc,
        ["Reference", "What it shows", "Key insight", "RQ / decision relevance"],
        [
            ["Figure 1", "National CPI and mandi trend", "CPI is smoother; mandi prices show large commodity shocks", "RQ1/RQ2: use separate targets and robust validation"],
            ["Figure 2", "Regional window rankings", "Pressure is persistent but horizon-sensitive", "RQ1/RQ4: avoid one-window policy rankings"],
            ["Coverage audit", "Reporting regions over time", "Panel is unbalanced", "All RQs: retain coverage controls"],
            ["Climate scatterplots", "Raw anomalies versus price change", "Near-zero bivariate correlation", "RQ1: require controlled lag models"],
        ],
        [1.0, 1.7, 1.9, 1.9],
        font_size=8.5,
    )
    add_heading(doc, "Choice of Models With Justification", 2)
    add_body(
        doc,
        "The model set follows the Step 4 design guidance while remaining proportionate to sample size. Ordinary least squares with heteroskedasticity-and-autocorrelation-consistent errors and state fixed-effects models address RQ1. One-step forecasting compares last-value and seasonal-naive baselines with Ridge autoregression with exogenous variables (ARX), Random Forest, and histogram gradient boosting. The 24-month task compares flat naive, seasonal naive, local drift, and damped exponential smoothing (ETS). Next-month shocks are modeled with persistence, balanced logistic regression, balanced Random Forest, and weighted histogram gradient boosting. Unsupervised and neural models were not added because they do not answer the defined targets or outperform the evidence standard merely by being more complex."
    )
    add_page_break(doc)

    # Page 11 - Features, metrics, performance
    add_heading(doc, "Features, Evaluation Metrics, and Preliminary Performance", 1)
    add_heading(doc, "Features Included and Feature Engineering", 2)
    add_table_caption(doc, 7, "Feature set and usage")
    add_table(
        doc,
        ["Feature family", "Original / engineered", "Reason", "Used in"],
        [
            ["Food CPI lags 1,2,3,6,12", "Engineered exact lags", "Persistence, seasonality, autoregression", "Ridge ARX and ML forecasts"],
            ["Rolling mean/std.", "Engineered from prior CPI", "Local level and volatility", "Forecast models"],
            ["Mandi index / YoY", "Original + engineered", "Wholesale pressure signal", "Regression, forecast, shock"],
            ["World wheat and crude", "Original + YoY change", "Global food and transport context", "Regression, forecast, shock"],
            ["Rainfall/temperature anomaly", "Engineered climate normal deviation", "Weather stress", "State regression and shock"],
            ["Month sine/cosine and trend", "Engineered", "Seasonality and long-run change", "Forecast and shock"],
        ],
        [1.6, 1.35, 1.9, 1.65],
        font_size=8.4,
    )
    add_heading(doc, "Evaluation Metrics (Formulae and Calculations)", 2)
    add_table_caption(doc, 8, "Evaluation metrics and formulae")
    add_table(
        doc,
        ["Metric", "Formula", "Interpretation", "Used for"],
        [
            ["MAE", "mean(|y - yhat|)", "Typical absolute index-point error", "Forecasting"],
            ["RMSE", "sqrt(mean((y - yhat)^2))", "Penalizes large misses", "Forecasting"],
            ["MAPE", "100 mean(|(y-yhat)/y|)", "Scale-free percentage error", "Forecasting"],
            ["Precision", "TP/(TP+FP)", "Share of alerts that are shocks", "Classification"],
            ["Recall", "TP/(TP+FN)", "Share of shocks detected", "Classification"],
            ["Balanced accuracy", "(Recall + Specificity)/2", "Balances rare shock and normal classes", "Classification"],
            ["Brier score", "mean((p-o)^2)", "Probability accuracy; lower is better", "Classification"],
            ["HFASI", "100 + w_food(g_food-g_power)", "Relative affordability stress", "Scenario index"],
        ],
        [1.15, 1.65, 2.25, 1.45],
        font_size=8.2,
    )
    add_body(
        doc,
        "The selected shock model produced TP=3, TN=52, FP=3, and FN=2. Therefore, precision=3/(3+3)=0.500, recall=3/(3+2)=0.600, specificity=52/(52+3)=0.945, and balanced accuracy=(0.600+0.945)/2=0.773. For the May 2026 rural downside scenario, HFASI=100+0.463(3.494-0)=101.618."
    )
    add_page_break(doc)

    # Page 12 - Model performance and RQ1/RQ2
    add_heading(doc, "Preliminary Results", 1)
    add_heading(doc, "Preliminary Model Performance", 2)
    add_table_caption(doc, 9, "Preliminary model performance comparison")
    add_table(
        doc,
        ["Task / model", "Validation", "Primary result", "Secondary result", "Key takeaway"],
        [
            ["One-step naive last", "60 rolling origins", "RMSE 2.063", "MAE 1.530; MAPE 1.071%", "Best one-step model"],
            ["One-step Ridge ARX", "60 rolling origins", "RMSE 2.572", "24.69% worse than naive", "Complexity did not improve accuracy"],
            ["24-step damped ETS", "24 origins x 24 horizons", "All-step RMSE 6.044", "H24 RMSE 8.391", "Best long-horizon path"],
            ["HistGradientBoosting shock", "60 rolling origins", "Balanced accuracy 0.773", "Recall 0.600; Brier 0.0758", "Best rare-event balance"],
        ],
        [1.45, 1.4, 1.25, 1.25, 1.65],
        font_size=8.4,
    )
    add_heading(doc, "RQ1 Preliminary Findings", 2)
    add_body(
        doc,
        "The national full model has adjusted R-squared=.870 and residual RMSE=1.243, but none of the five pre-specified national transmission terms is significant at 5%. State fixed-effects models identify a positive lagged temperature-anomaly association in the two-year (coefficient 4.99, Holm-adjusted p<.001) and five-year (4.28, Holm-adjusted p<.001) windows, but not the ten-year window. Rainfall terms are not significant after Holm adjustment. RQ1 therefore receives mixed, horizon-dependent associational support, not causal confirmation."
    )
    add_heading(doc, "RQ2 Preliminary Findings", 2)
    add_body(
        doc,
        "Last-value persistence wins the one-step comparison, whereas damped ETS wins the genuine 24-step backtest. This rejects the assumption that a machine-learning model must be operationally superior. Ridge permutation importance is highest for the 12-month CPI lag (mean RMSE degradation 1.01), followed by the one-month lag and mandi index. The 2026 point path implies average YoY food-cost growth of 2.58%, while empirical intervals widen materially with horizon."
    )
    add_figure_caption(doc, 3, "Twenty-four-month conditional national food CPI forecast")
    add_picture(doc, ASSETS / "nb05_fig3.png", width=6.15)
    add_note(doc, "The damped ETS point path is selected by 24-step backtesting. Empirical intervals use matching historical horizons and are approximate because only 24 origins are available.")
    add_page_break(doc)

    # Page 13 - RQ3/RQ4
    add_heading(doc, "Preliminary Findings by Research Question (Continued)", 1)
    add_heading(doc, "RQ3 Preliminary Findings", 2)
    add_body(
        doc,
        "The shock threshold is the initial-training 75th percentile of next-month YoY food inflation (8.591%). HistGradientBoosting achieves 0.500 precision, 0.600 recall, 0.945 specificity, ROC-AUC 0.764, average precision 0.452, and Brier score 0.0758. Only five shocks occur among 60 evaluation months, so the estimates remain uncertain. The January 2026 probability based entirely on observed inputs is 2.59%; February-December are conditional on the damped ETS CPI path, constant December 2025 external-price levels, and seasonal-normal climate. The maximum baseline probability is 4.12% in August."
    )
    add_body(
        doc,
        "HFASI uses published food shares of 0.463 rural and 0.391 urban. With 4% purchasing-power growth, the highest baseline HFASI is 99.80, so point estimates imply relative improvement rather than stress. With 0% purchasing-power growth, all 12 months show mild stress; rural mean HFASI is 101.20 versus 101.01 urban. Forecast uncertainty crosses 100 in multiple months, so the baseline point conclusion is not guaranteed and cannot be interpreted as absolute household prosperity."
    )
    add_heading(doc, "RQ4 Preliminary Findings", 2)
    add_body(
        doc,
        "The decision layer combines maximum monthly HFASI with shock probability. Baseline and favorable-supply scenarios remain normal. Moderate stress produces an HFASI budget watch but no procurement alert. Severe stress produces elevated household stress and a procurement alert: rural HFASI reaches 102.54, maximum shock probability reaches 74.88%, nine months exceed a 10% watch level, and three exceed the 50% classifier cutoff. These thresholds demonstrate a transparent workflow; they are not optimized policy rules."
    )
    add_figure_caption(doc, 4, "Conditional shock probability under alternative 2026 scenarios")
    add_picture(doc, ASSETS / "nb09_fig1.png", width=6.25)
    add_note(doc, "The severe scenario is a controlled stress test, not a forecast. The favorable scenario's slightly higher mean probability than baseline reflects nonlinear model sensitivity and must not be given a causal interpretation.")
    add_page_break(doc)

    # Page 14 - limitations and next steps
    add_heading(doc, "Interim Limitations and Risks", 1)
    add_body(
        doc,
        "Household representation is the principal scope limitation. HFASI uses published aggregate rural and urban food shares because the HCES microdata were not supplied. It cannot support household-level, expenditure-quantile, or causal welfare inference. The affordability results are representative scenarios whose purchasing-power assumptions must remain visible."
    )
    add_body(
        doc,
        "Rare-event uncertainty is also material. The 60-month shock evaluation contains only five shocks, so recall, precision, and probability estimates have high sampling uncertainty. The classifier is useful as a risk watch but not as an automatic policy trigger. The favorable scenario's slightly higher mean risk than baseline further demonstrates nonlinear model sensitivity."
    )
    add_body(
        doc,
        "The AGMARKNET panel is unbalanced, coverage changes can resemble price changes, and the state regressions can contain omitted variables, common trends, or release-timing mismatches. Climate associations are horizon-dependent and not causal. Long-horizon intervals use only 24 historical forecast origins. These issues are mitigated through coverage indicators, exact calendar joins, raw-value preservation, cluster/HAC uncertainty, chronological validation, sensitivity analysis, and explicit interpretation constraints."
    )
    add_heading(doc, "Next Steps for the Final Report", 1)
    add_bullet(doc, "Obtain authorized HCES 2022-23/2023-24 microdata, apply survey weights, and validate HFASI across expenditure fractiles without publishing restricted raw records.")
    add_bullet(doc, "Expand rare-shock evaluation as new observed months become available; add calibrated alert thresholds based on asymmetric decision costs.")
    add_bullet(doc, "Backtest household-budget and procurement losses against no-warning, persistence, and seasonal-naive decision rules.")
    add_bullet(doc, "Evaluate additional official wage, arrival, production, and policy-event features only when release dates prevent leakage.")
    add_bullet(doc, "Complete model cards, environment lock files, automated tests, and monitored drift/coverage checks before deployment.")
    add_heading(doc, "Interim Conclusion", 1)
    add_body(
        doc,
        "The interim evidence supports an integrated but deliberately conservative framework. Strong baselines remain essential; predictive drivers are not causal; affordability conclusions depend on household exposure and purchasing power; and severe scenarios can jointly trigger household and procurement warnings. The completed pipeline is reproducible, explainable, and suitable for final validation, while its limitations are sufficiently material to prevent overstatement."
    )
    add_page_break(doc)

    # Pages 15-16 - bibliography
    add_heading(doc, "Bibliography", 1)
    references = [
        "Akter, S., & Basher, S. A. (2014). The impacts of food price and income shocks on household food security and economic well-being: Evidence from rural Bangladesh. Global Environmental Change, 25, 150-162. https://doi.org/10.1016/j.gloenvcha.2014.02.003",
        "Bhardwaj, M. R., Pawar, J., Bhat, A., Deepanshu, Enaganti, I., Sagar, K., & Narahari, Y. (2023). An innovative deep learning based approach for accurate agricultural crop price prediction. In 2023 IEEE 19th International Conference on Automation Science and Engineering. https://doi.org/10.1109/CASE56687.2023.10260494",
        "Cattaneo, A., Sadiddin, A., Vaz, S., Conti, V., Holleman, C., Sanchez, M. V., & Torero, M. (2023). Ensuring affordability of diets in the face of shocks. Food Policy, 117, 102470. https://doi.org/10.1016/j.foodpol.2023.102470",
        "Hyndman, R. J., & Koehler, A. B. (2006). Another look at measures of forecast accuracy. International Journal of Forecasting, 22(4), 679-688. https://doi.org/10.1016/j.ijforecast.2006.03.001",
        "Jain, A., Marvaniya, S., Godbole, S., & Munigala, V. (2020). A framework for crop price forecasting in emerging economies by analyzing the quality of time-series data. arXiv. https://doi.org/10.48550/arXiv.2009.04171",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30, 4765-4774.",
        "Ma, W., Nowocin, K., Marathe, N., & Chen, G. H. (2019). An interpretable produce price forecasting system for small and marginal farmers in India using collaborative filtering and adaptive nearest neighbors. Proceedings of the Tenth International Conference on Information and Communication Technologies and Development, Article 6, 1-11. https://doi.org/10.1145/3287098.3287100",
        "Madaan, L., Sharma, A., Khandelwal, P., Goel, S., Singla, P., & Seth, A. (2019). Price forecasting and anomaly detection for agricultural commodities in India. Proceedings of the 2nd ACM SIGCAS Conference on Computing and Sustainable Societies. https://doi.org/10.1145/3314344.3332488",
        "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2018). Statistical and machine learning forecasting methods: Concerns and ways forward. PLOS ONE, 13(3), e0194889. https://doi.org/10.1371/journal.pone.0194889",
        "Ministry of Statistics and Programme Implementation. (2024). Survey on Household Consumption Expenditure: 2022-23 (NSS Report No. 591). Government of India. https://mospi.gov.in/sites/default/files/publication_reports/Report_591_HCES_2022-23New.pdf",
        "Open Government Data Platform India. (2026). Current daily price of various commodities from various markets (Mandi). https://www.data.gov.in/catalog/current-daily-price-various-commodities-various-markets-mandi",
        "Theofilou, A., Arvanitakis, G., & Tzouramani, I. (2025). Predicting prices of staple crops using machine learning: A systematic review of studies on wheat, corn, and rice. Sustainability, 17(12), 5456. https://doi.org/10.3390/su17125456",
        "World Bank. (2026). Commodity markets (Pink Sheet) data. https://www.worldbank.org/en/research/commodity-markets",
        "National Aeronautics and Space Administration. (2026). POWER Data Access Viewer. https://power.larc.nasa.gov/",
    ]
    for ref in references:
        add_reference(doc, ref)
    add_page_break(doc)

    # Appendix A
    add_heading(doc, "Appendix A", 1)
    add_heading(doc, "GitHub Repository Structure and Reproducibility Map", 2)
    add_body(
        doc,
        "The repository structure below is referenced in the GitHub availability statement and maps evaluator actions to stable public paths."
    )
    tree = [
        "qm640-food-affordability-ai/",
        "|-- README.md",
        "|-- config/                 # versioned settings and source profiles",
        "|-- data/",
        "|   |-- curated/            # evaluator-accessible compressed datasets",
        "|   |-- raw/                # local-only bulk or credential-dependent data",
        "|   |-- data_dictionary.csv",
        "|   `-- source_manifest.csv",
        "|-- notebooks/",
        "|   |-- 01_data_acquisition.ipynb",
        "|   |-- 02_data_quality_and_cleaning.ipynb",
        "|   |-- 03_exploratory_analysis.ipynb",
        "|   |-- 04_statistical_analysis.ipynb",
        "|   |-- 05_forecasting_models.ipynb",
        "|   |-- 06_shock_classification.ipynb",
        "|   |-- 07_explainability.ipynb",
        "|   |-- 08_affordability_index.ipynb",
        "|   |-- 09_scenario_analysis.ipynb",
        "|   `-- 10_final_results.ipynb",
        "|-- scripts/                # collection, build, validation, report code",
        "|-- src/                    # reusable Python package",
        "|-- tests/                  # automated quality and pipeline tests",
        "`-- reports/                # figures, tables, model cards, final report",
    ]
    p = doc.add_paragraph()
    for line in tree:
        r = p.add_run(line + "\n")
        r.font.name = "Courier New"
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Courier New")
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Courier New")
        r.font.size = Pt(8.5)
    format_paragraph(
        p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT,
        double=False, after=4,
    )
    add_table_caption(doc, 10, "Notebook purpose and primary verified output")
    add_table(
        doc,
        ["Notebook", "Purpose", "Primary output"],
        [
            ["01", "Acquisition and panel construction", "14 curated files; 1,494,745 rows"],
            ["02", "Quality and cleaning", "41,792 clean state rows; no price interpolation"],
            ["03", "EDA", "2/5/10-year pressure, volatility, coverage"],
            ["04", "Statistical analysis", "HAC national and fixed-effects state results"],
            ["05", "Forecasting", "Naive one-step; damped ETS 24-step path"],
            ["06", "Shock classification", "HGB metrics and 2026 monthly risk"],
            ["07", "Explainability", "Permutation, partial dependence, local sensitivity"],
            ["08", "HFASI", "Rural/urban scenario and uncertainty tables"],
            ["09", "Scenario decisions", "Baseline through severe stress signals"],
            ["10", "Consolidation", "RQ map, dashboard, limitations, 98-artifact inventory"],
        ],
        [0.75, 2.7, 3.05],
        font_size=8.2,
    )
    add_page_break(doc)

    # Appendix B
    add_heading(doc, "Appendix B", 1)
    add_heading(doc, "Extended Formulae, Assumptions, and Result Checks", 2)
    formulas = [
        ["MAE", "(1/n) SUM |y_t - yhat_t|", "Naive one-step MAE = 1.5295"],
        ["RMSE", "sqrt[(1/n) SUM (y_t-yhat_t)^2]", "Naive one-step RMSE = 2.0628"],
        ["MAPE", "(100/n) SUM |(y_t-yhat_t)/y_t|", "Naive MAPE = 1.0712%"],
        ["Precision", "TP/(TP+FP)", "3/(3+3) = 0.500"],
        ["Recall", "TP/(TP+FN)", "3/(3+2) = 0.600"],
        ["Specificity", "TN/(TN+FP)", "52/(52+3) = 0.945"],
        ["Balanced accuracy", "(Recall+Specificity)/2", "(0.600+0.945)/2 = 0.773"],
        ["Brier", "(1/n) SUM (p_t-o_t)^2", "Selected classifier = 0.0758"],
        ["HFASI", "100+w_food(g_food-g_power)", "Rural May downside = 101.618"],
    ]
    add_table_caption(doc, 11, "Formulae and actual interim calculations")
    add_table(doc, ["Measure", "Formula", "Verified calculation"], formulas, [1.3, 2.9, 2.3], font_size=8.5)
    add_heading(doc, "Conditional 2026 Assumptions", 2)
    add_bullet(doc, "January shock risk uses observed December 2025 inputs.")
    add_bullet(doc, "February-December shock risks use the Notebook 05 CPI path.")
    add_bullet(doc, "External price levels are held at December 2025 values.")
    add_bullet(doc, "Future climate anomalies are set to seasonal normal in the baseline.")
    add_bullet(doc, "HFASI purchasing-power scenarios are 0%, 4%, and 8%; they are assumptions, not income forecasts.")
    add_heading(doc, "Stress-Scenario Overlays", 2)
    add_table(
        doc,
        ["Scenario", "Food-cost overlay", "Mandi / wheat / crude", "Rain / temperature"],
        [
            ["Baseline", "0 pp", "0 / 0 / 0 pp", "0 pp / 0 C"],
            ["Favorable supply", "-2 pp", "-5 / -5 / -10 pp", "+10 pp / -0.5 C"],
            ["Moderate stress", "+3 pp", "+10 / +10 / +15 pp", "-15 pp / +1.5 C"],
            ["Severe stress", "+6 pp", "+20 / +20 / +30 pp", "-30 pp / +3.0 C"],
        ],
        [1.45, 1.25, 2.1, 1.7],
        font_size=8.5,
    )
    add_note(doc, "Scenario changes are controlled feature perturbations. They are not probabilistic forecasts or estimated causal effects.")
    add_page_break(doc)

    # Appendix C
    add_heading(doc, "Appendix C", 1)
    add_heading(doc, "Consolidated Results Dashboard and Artifact Evidence", 2)
    add_body(
        doc,
        "Figure C1 consolidates the principal forecast, baseline risk, affordability, and scenario decision outputs. It is generated by Notebook 10 and is referenced in the preliminary-results discussion."
    )
    add_figure_caption(doc, 5, "Food Price Affordability AI consolidated interim results")
    add_picture(doc, ASSETS / "nb10_fig1.png", width=6.45)
    add_note(doc, "Panels show the 24-month damped ETS path, baseline 2026 shock risk, baseline rural/urban HFASI, and the joint scenario decision space. All panels are conditional and non-causal.")
    add_heading(doc, "Verified Final Evidence", 2)
    add_bullet(doc, "Notebooks 01-09 completed and verified by Notebook 10.")
    add_bullet(doc, "Four research questions mapped to evidence and interpretation strength.")
    add_bullet(doc, "Eleven key findings and six material limitations consolidated.")
    add_bullet(doc, "Ninety-eight report artifacts inventoried with file size and SHA-256 hash.")
    add_bullet(doc, "No missing prices interpolated; raw extreme values preserved with flags.")
    add_bullet(doc, "Report numbers are transcribed from execution summaries, not manually recalculated.")
    return doc


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    for required_asset in [
        "nb03_fig1.png", "nb03_fig4.png", "nb05_fig3.png",
        "nb09_fig1.png", "nb10_fig1.png",
    ]:
        if not (ASSETS / required_asset).exists():
            raise FileNotFoundError(ASSETS / required_asset)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.core_properties.title = (
        "QM640 Interim Report - Food Price Affordability AI"
    )
    doc.core_properties.subject = "Data Analytics Capstone Interim Report"
    doc.core_properties.author = "Piyush Soni"
    doc.core_properties.keywords = (
        "food price, affordability, forecasting, shock classification, India"
    )
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
